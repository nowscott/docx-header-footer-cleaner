from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import shutil
import tempfile

def clear_hf(hf):
    el = hf._element
    for child in list(el):
        el.remove(child)

def add_center_page_number(footer):
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()._r
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)

def process_document(input_path, output_path):
    doc = Document(input_path)
    for section in doc.sections:
        try:
            section.header.is_linked_to_previous = False
        except Exception:
            pass
        try:
            section.footer.is_linked_to_previous = False
        except Exception:
            pass
        try:
            section.different_first_page_header_footer = False
        except Exception:
            pass
        clear_hf(section.header)
        clear_hf(section.footer)
        try:
            clear_hf(section.first_page_header)
            clear_hf(section.first_page_footer)
        except Exception:
            pass
        try:
            clear_hf(section.even_page_header)
            clear_hf(section.even_page_footer)
        except Exception:
            pass
        add_center_page_number(section.footer)
    doc.save(output_path)

def derive_output_path(in_path, out_path):
    if out_path:
        return out_path
    base, ext = os.path.splitext(in_path)
    return base + '_paged' + ext

def read_config(path):
    roots = []
    cfg = {}
    with open(path, 'r', encoding='utf-8') as f:
        for line in f:
            s = line.strip()
            if not s:
                continue
            if s.startswith('#'):
                continue
            if '=' in s:
                k, v = s.split('=', 1)
                cfg[k.strip().lower()] = v.strip()
                continue
            roots.append(s)
    return roots, cfg

def ensure_dir(p):
    os.makedirs(p, exist_ok=True)

def backup_file(src, root, backup_root):
    rel = os.path.relpath(src, root)
    dest_root = os.path.join(backup_root, os.path.basename(root))
    dest_path = os.path.join(dest_root, rel)
    ensure_dir(os.path.dirname(dest_path))
    shutil.copy2(src, dest_path)
    return dest_path

def process_in_place(path):
    d = os.path.dirname(path)
    fd, tmp = tempfile.mkstemp(prefix='__paged_', suffix='.docx', dir=d)
    os.close(fd)
    try:
        process_document(path, tmp)
        os.replace(tmp, path)
    finally:
        if os.path.exists(tmp):
            try:
                os.remove(tmp)
            except Exception:
                pass

def process_roots(roots, backup_root):
    processed = []
    errors = []
    skipped = []
    for root in roots:
        for dirpath, dirnames, filenames in os.walk(root):
            for fn in filenames:
                lower = fn.lower()
                full = os.path.join(dirpath, fn)
                if lower.endswith('.docx'):
                    try:
                        backup_file(full, root, backup_root)
                        process_in_place(full)
                        processed.append(full)
                    except Exception as e:
                        errors.append((full, str(e)))
                elif lower.endswith('.doc'):
                    skipped.append(full)
    return processed, errors, skipped

def main():
    import argparse
    import sys
    parser = argparse.ArgumentParser(description='移除页眉页脚并在页脚居中添加页码')
    parser.add_argument('input', nargs='?', help='输入的.docx文件路径')
    parser.add_argument('-o', '--output', help='输出文件路径')
    parser.add_argument('--config', help='配置文件路径，包含备份目录和待处理目录')
    parser.add_argument('--backup', help='备份原文件的目录，优先于配置文件设置')
    args = parser.parse_args()
    default_cfg = os.path.join(os.getcwd(), 'docx_config.txt')
    local_cfg = os.path.join(os.getcwd(), 'docx_config.local.txt')
    cfg_path = args.config if args.config else (local_cfg if os.path.isfile(local_cfg) else default_cfg)
    if os.path.isfile(cfg_path):
        roots, cfg = read_config(cfg_path)
        if not roots:
            print('配置文件无有效目录')
            sys.exit(1)
        backup_root = args.backup or cfg.get('backup') or os.path.join(os.getcwd(), 'docx_backup')
        ensure_dir(backup_root)
        processed, errors, skipped = process_roots(roots, backup_root)
        print('SUMMARY')
        print('Processed .docx:', len(processed))
        print('Errors:', len(errors))
        print('Skipped .doc (convert to .docx first):', len(skipped))
        if processed:
            print('\nDETAILS (first 30):')
            for i, p in enumerate(processed[:30], 1):
                print(f'{i:02d}.', p)
        if errors:
            print('\nERRORS (first 20):')
            for i, (inp, err) in enumerate(errors[:20], 1):
                print(f'{i:02d}.', inp, '\n   ', err)
        if skipped:
            print('\nSKIPPED .doc (first 20):')
            for i, p in enumerate(skipped[:20], 1):
                print(f'{i:02d}.', p)
        return
    if not args.input:
        print('缺少输入文件或--config')
        sys.exit(1)
    if not os.path.isfile(args.input):
        print('输入文件不存在')
        sys.exit(1)
    if not args.input.lower().endswith('.docx'):
        print('仅支持.docx格式')
        sys.exit(1)
    out_path = derive_output_path(args.input, args.output)
    process_document(args.input, out_path)
    print(out_path)

if __name__ == '__main__':
    main()